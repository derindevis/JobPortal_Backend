from fastapi import APIRouter, HTTPException, Depends, status, File, UploadFile, Form
import PyPDF2
import os
import uuid
import docx
from io import BytesIO
from utils.ai_evaluator import evaluate_resume
from sqlalchemy.orm import Session
from typing import List
from database import get_db
from models.application import Application
from schemas.application import ApplicationCreate, ApplicationResponse, ApplicationStatusUpdate
from dependencies import get_current_user, require_admin, require_user
from models.job import Job
from models.user import User

router=APIRouter()

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

@router.get("/me", response_model=List[ApplicationResponse])
def get_dashboard(db: Session=Depends(get_db), current_user=Depends(get_current_user)):
    if current_user.role=="admin":
        return db.query(Application).all()
    return db.query(Application).filter(Application.user_id==current_user.id).all()
    

@router.post("/", response_model=ApplicationResponse, status_code=201)
def apply_for_job(
    job_id: int = Form(...),
    cover_letter: str = Form(...),
    resume: UploadFile = File(None),
    db: Session=Depends(get_db), 
    current_user:User=Depends(require_user)
):
    if current_user.role == "admin":
        raise HTTPException(status_code=403, detail="Admins cannot apply for jobs.")
        
    job=db.query(Job).filter(Job.id==job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="The job doesn't exist.")
    if not job.active:
        raise HTTPException(status_code=400, detail="This job is no longer accepting applications.")
    
    existing=db.query(Application).filter(Application.job_id==job_id, Application.user_id==current_user.id).first()
    if existing:
        raise HTTPException(status_code=400, detail="You have already applied!")

    # Parse resume text if uploaded and save the file
    extracted_text = ""
    resume_path = None
    file_processing_error = None
    if resume and resume.filename:
        try:
            # Generate unique filename
            ext = os.path.splitext(resume.filename)[1]
            unique_filename = f"{uuid.uuid4()}{ext}"
            file_path = os.path.join(UPLOAD_DIR, unique_filename)
            
            # Read file bytes (limit to 5MB to prevent memory exhaustion)
            file_bytes = resume.file.read(5 * 1024 * 1024)
            
            # Save file to uploads folder
            with open(file_path, "wb") as f:
                f.write(file_bytes)
            
            resume_path = f"uploads/{unique_filename}"
            
            if resume.filename.endswith('.pdf'):
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
                for page in pdf_reader.pages:
                    extracted_text += page.extract_text() + "\n"
            elif resume.filename.endswith('.docx'):
                doc = docx.Document(BytesIO(file_bytes))
                fullText = []
                for para in doc.paragraphs:
                    fullText.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            fullText.append(cell.text)
                extracted_text = '\n'.join(fullText)
        except Exception as e:
            print(f"Failed to process uploaded file: {e}")
            file_processing_error = str(e)

    # AI Evaluation
    eval_result = evaluate_resume(
        resume_text=extracted_text if extracted_text.strip() else cover_letter,
        job_description=job.description,
        cover_letter=cover_letter
    )
    
    ai_reasoning = eval_result.get("reasoning")
    if file_processing_error:
        ai_reasoning = f"[File Processing Error: {file_processing_error}] {ai_reasoning or ''}"

    new_app=Application(
        job_id=job_id, 
        user_id=current_user.id, 
        cover_letter=cover_letter,
        ai_score=eval_result.get("score"),
        ai_reasoning=ai_reasoning,
        resume_path=resume_path
    )
    db.add(new_app)
    db.commit()
    db.refresh(new_app)
    return new_app


@router.put("/{app_id}", response_model=ApplicationResponse)
def update_app_status(app_id:int, payload: ApplicationStatusUpdate, db: Session=Depends(get_db),admin=Depends(require_admin)):
    application=db.query(Application).filter(Application.id==app_id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application Not Found!")
    application.status=payload.status
    db.commit()
    db.refresh(application)
    return application


@router.delete("/{app_id}")
def withdraw_application(
    app_id: int,
    db: Session = Depends(get_db),
    current_user=Depends(get_current_user)
):
    application=db.query(Application).filter(Application.id==app_id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application Not Found")
    
    if application.user_id!=current_user.id and current_user.role != "admin":
        raise HTTPException(status_code=403, detail="You are not authorized to withdraw this application.")

    db.delete(application)
    db.commit()
    return {"Message":"Application Withdrawn Successfully!"}